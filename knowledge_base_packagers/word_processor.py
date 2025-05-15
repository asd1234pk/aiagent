'''
Module to process Word documents from a local directory.
'''
import os
from docx import Document
import re

# --- Configuration ---
# 使用者可以修改此路徑以指向其實際的 Word 文件資料夾
# 建議使用絕對路徑或相對於專案根目錄的路徑
WORD_DOCS_DIR = "knowledge_docs/word_documents" 

# 針對 FAQ 和說明型文件，嘗試較小的區塊以增強每個獨立問答的獨特性
CHUNK_SIZE_CHARS = 500 # 原為 1000
CHUNK_OVERLAP_CHARS = 50  # 原為 100
# --- End Configuration ---

def ensure_docs_dir_exists():
    """檢查 WORD_DOCS_DIR 是否存在，如果不存在則建立它。"""
    if not os.path.exists(WORD_DOCS_DIR):
        print(f"指定的 Word 文件目錄 '{WORD_DOCS_DIR}' 不存在。")
        try:
            os.makedirs(WORD_DOCS_DIR)
            print(f"已建立目錄：'{WORD_DOCS_DIR}'。請將您的 Word 文件放入此目錄中再執行同步。")
            # 可以在此處放置一個範例檔案供使用者參考
            # example_doc_path = os.path.join(WORD_DOCS_DIR, "example_document.docx")
            # if not os.path.exists(example_doc_path):
            #     doc = Document()
            #     doc.add_paragraph("這是一個範例 Word 文件，用於測試知識庫 Word 文件處理功能。")
            #     doc.add_paragraph("您可以刪除此檔案，並將您自己的 .docx 文件放入此目錄中。")
            #     doc.save(example_doc_path)
            #     print(f"已在 '{WORD_DOCS_DIR}' 中建立一個範例文件: example_document.docx")
        except OSError as e:
            print(f"建立目錄 '{WORD_DOCS_DIR}' 時發生錯誤: {e}")
            print("請手動建立此目錄並放入 Word 文件，或檢查權限。")
            return False
    return True

def extract_text_from_docx(file_path):
    """從單個 .docx 文件中提取所有文本。"""
    try:
        doc = Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            # 考慮是否需要處理表格、頁首頁尾等，目前僅處理段落
            # 可以透過 para.runs 來獲取更細緻的格式資訊，但純文字提取通常 paragraph.text 已足夠
            # 移除段落前後可能存在的額外空白
            paragraph_text = para.text.strip()
            if paragraph_text: # 只加入非空段落
                full_text.append(paragraph_text)
        return "\n".join(full_text) # 以換行符連接段落
    except Exception as e:
        print(f"讀取 Word 文件 '{file_path}' 時發生錯誤: {e}")
        return None

def split_text_into_chunks(text, source_file_path, document_title):
    """將長文本切分成較小的區塊。"""
    if not text:
        return []
    
    chunks = []
    # 簡易的純文字切分，可以考慮更複雜的句子邊界檢測
    # 這裡的實現與 website_scraper 中的基本一致，只是 source 和 title 的意義不同
    current_position = 0
    text_len = len(text)

    while current_position < text_len:
        end_position = min(current_position + CHUNK_SIZE_CHARS, text_len)
        # 先暫存一個基礎區塊，如果找不到更好的斷點就用它
        current_chunk_text = text[current_position:end_position]
        final_end_position = end_position
        
        # 嘗試向後找到自然的斷點 (如句子結束符或多個換行)
        if end_position < text_len: # 如果不是最後一個區塊
            # 搜尋範圍稍微擴大一點，以便包含可能的標點或換行
            search_look_ahead = int(CHUNK_OVERLAP_CHARS * 1.5) # 比重疊稍大，給予尋找斷點的空間
            search_end = min(end_position + search_look_ahead, text_len)
            
            # 優先考慮雙換行 (通常代表段落結束或不同 FAQ 條目間隔)
            # regex: 匹配點/問/嘆號/句號/問號/嘆號 + 空白，或者一個或多個換行符
            # 將 \n+ (多個換行) 放在前面，使其優先級高於單個 \n

            # 尋找斷點的 regex: [.!?。？！]\s+ (句尾標點加空格) OR \n{2,} (兩個以上換行) OR \n (單個換行，作為次選)
            # 注意: re.finditer 會找到不重疊的匹配。我們是從 end_position 開始向後搜索。 
            # 先嘗試匹配句尾標點或多個換行符
            primary_split_regex = r'[.!?。？！]\s+|'
            sentence_boundaries = []
            for match in re.finditer(primary_split_regex, text[current_position:search_end]):
                # 我們關心的是在 CHUNK_SIZE_CHARS 附近的斷點
                # match.start() 是相對於 text[current_position:search_end] 的位置
                # 所以實際的斷點位置是 current_position + match.start() + len(match.group(0))
                potential_split_at = match.start() + len(match.group(0))
                if current_position + potential_split_at > current_position + CHUNK_SIZE_CHARS - int(CHUNK_SIZE_CHARS * 0.3): # 在區塊後70%部分找到的斷點才考慮
                    if current_position + potential_split_at <= search_end: # 確保斷點在搜索範圍內
                        sentence_boundaries.append(potential_split_at)
                        break # 找到一個就用，優先用靠前的強斷點
            if sentence_boundaries: 
                final_end_position = current_position + sentence_boundaries[0]
                current_chunk_text = text[current_position:final_end_position]
            else:
                # 如果沒有找到主要斷點，再嘗試單個換行符 (如果 CHUNK_SIZE_CHARS 允許)
                # 這種情況通常是希望在長段落中根據換行做一些切分
                secondary_split_regex = r'\n'
                line_boundaries = []
                for match in re.finditer(secondary_split_regex, text[current_position:search_end]):
                    potential_split_at = match.start() + len(match.group(0))
                    if current_position + potential_split_at > current_position + CHUNK_SIZE_CHARS - int(CHUNK_SIZE_CHARS * 0.3):
                        if current_position + potential_split_at <= search_end: 
                            line_boundaries.append(potential_split_at)
                            break # 找到一個就用
                if line_boundaries:
                    final_end_position = current_position + line_boundaries[0]
                    current_chunk_text = text[current_position:final_end_position]
                # 若都找不到，則使用原始的基于 CHUNK_SIZE_CHARS 的切分 (即 current_chunk_text 和 final_end_position 已是預設)

        chunks.append({
            "text": current_chunk_text.strip(), # 確保移除區塊前後的空白
            "source": source_file_path, # Word 檔案的完整路徑
            "title": document_title, # Word 檔案的名稱 (不含副檔名)
            "type": "word_document", # 標示此區塊來自 Word 文件
            # 可以考慮加入 chunk_index 或其他元數據
            "chunk_char_start": current_position,
            "chunk_char_end": final_end_position
        })
        
        if final_end_position == text_len:
            break # 已到達文本末尾
            
        # 移動到下一個區塊的起始位置，考慮重疊
        current_position = max(current_position + 1, final_end_position - CHUNK_OVERLAP_CHARS)
        if current_position >= final_end_position: # 避免死循環
            current_position = final_end_position
            
    return [c for c in chunks if c["text"]] # 過濾掉可能產生的空文本區塊

def process_all_documents():
    """
    掃描 WORD_DOCS_DIR 目錄，處理所有 .docx 文件，提取文本並切分成區塊。
    返回一個包含所有文本區塊的列表。
    """
    if not ensure_docs_dir_exists():
        return [] # 如果目錄處理失敗，返回空列表

    all_text_chunks = []
    file_count = 0
    processed_file_count = 0

    print(f"開始掃描目錄 '{WORD_DOCS_DIR}' 中的 Word 文件 (.docx)...")

    for root, _, files in os.walk(WORD_DOCS_DIR):
        for filename in files:
            if filename.lower().endswith(".docx") and not filename.startswith("~"): # 忽略 Word 暫存檔
                file_count += 1
                file_path = os.path.join(root, filename)
                document_title = os.path.splitext(filename)[0] # 取檔案名作為標題
                print(f"  處理中文件: {file_path}")
                
                raw_text = extract_text_from_docx(file_path)
                if raw_text:
                    chunks = split_text_into_chunks(raw_text, file_path, document_title)
                    if chunks:
                        all_text_chunks.extend(chunks)
                        print(f"    -> 從 '{filename}' 提取並切分了 {len(chunks)} 個文本區塊。")
                        processed_file_count += 1
                    else:
                        print(f"    -> 從 '{filename}' 未能切分出任何文本區塊 (原始文本可能為空或過短)。")
                else:
                    print(f"    -> 從 '{filename}' 提取文本失敗。")
    
    print(f"Word 文件掃描完成。共找到 {file_count} 個 .docx 文件，成功處理了 {processed_file_count} 個。")
    print(f"總共提取了 {len(all_text_chunks)} 個文本區塊。")
    return all_text_chunks

# --- Main execution for testing ---
if __name__ == '__main__':
    print("--- 測試 Word 文件處理器 ---")
    # 執行前，請確保 WORD_DOCS_DIR 指向一個包含 .docx 文件的目錄
    # 或者該目錄存在且為空（腳本會嘗試建立它並提示放入文件）

    # 檢查/建立目錄
    # ensure_docs_dir_exists() # process_all_documents 會呼叫它

    processed_chunks = process_all_documents()

    if processed_chunks:
        print(f"\n成功從 Word 文件處理了 {len(processed_chunks)} 個文本區塊。")
        print("顯示前 3 個區塊的範例：")
        for i, chunk in enumerate(processed_chunks[:min(3, len(processed_chunks))]):
            print(f"--- 區塊 {i+1} ---")
            print(f"  標題: {chunk['title']}")
            print(f"  來源: {chunk['source']}")
            # print(f"  類型: {chunk['type']}")
            # print(f"  起始: {chunk['chunk_char_start']}, 結束: {chunk['chunk_char_end']}")
            print(f"  文本: {chunk['text'][:200].replace(chr(10), ' ')}...") # 替換換行符以利於單行預覽
    else:
        print("\n未能從 Word 文件處理任何文本區塊。")
        print(f"請檢查 '{WORD_DOCS_DIR}' 目錄中是否有 .docx 文件，以及是否有讀取權限。")
    
    print("--- Word 文件處理器測試完成 ---") 